import os
import logging
import json
import hashlib
from typing import List, Dict, Any, Optional
from pathlib import Path
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument
from .config import settings
from .database import get_db_session
from .models import Chunk

logger = logging.getLogger(__name__)


def calculate_file_hash(file_path: str) -> str:
    """
    Вычисляет SHA256 хэш-сумму файла
    
    Args:
        file_path: Путь к файлу
        
    Returns:
        SHA256 хэш-сумма в виде hex-строки
    """
    sha256_hash = hashlib.sha256()
    try:
        with open(file_path, "rb") as f:
            # Читаем файл по частям для экономии памяти
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    except Exception as e:
        logger.error(f"Ошибка при вычислении хэша файла {file_path}: {e}")
        return ""


def load_docx_files(docs_path: str) -> List[Dict[str, Any]]:
    """
    Загружает все .docx файлы из указанной папки
    """
    documents = []
    docs_dir = Path(docs_path)
    
    if not docs_dir.exists():
        logger.warning(f"Папка документов не найдена: {docs_path}")
        return documents
    
    for file_path in docs_dir.glob("*.docx"):
        try:
            logger.info(f"Загрузка документа: {file_path}")
            text = extract_text_from_docx(str(file_path))
            if text.strip():
                file_hash = calculate_file_hash(str(file_path))
                documents.append({
                    "title": file_path.stem,
                    "content": text,
                    "file_path": str(file_path),
                    "file_size": file_path.stat().st_size,
                    "file_hash": file_hash
                })
                logger.info(f"Документ загружен: {file_path.name} ({len(text)} символов, hash: {file_hash[:16]}...)")
            else:
                logger.warning(f"Документ пустой: {file_path}")
        except Exception as e:
            logger.error(f"Ошибка при загрузке документа {file_path}: {e}")
    
    logger.info(f"Загружено документов: {len(documents)}")
    return documents


def extract_text_from_docx(file_path: str) -> str:
    """
    Извлекает текст из .docx файла
    """
    try:
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Ошибка при извлечении текста из {file_path}: {e}")
        raise


def save_chunks_to_db(chunks: List[LangChainDocument]) -> List[int]:
    """
    Сохраняет чанки в базу данных
    
    Перед сохранением новых чанков обновляет статус всех существующих чанков на "stored".
    Новые чанки сохраняются со статусом "actual".
    
    Args:
        chunks: Список LangChain документов (чанков)
        
    Returns:
        Список ID сохраненных чанков
    """
    chunk_ids = []
    try:
        with get_db_session() as db:
            # Перед сохранением новых чанков помечаем все существующие как "stored"
            from sqlalchemy import text
            result = db.execute(
                text('UPDATE "oozo-schema".chunks SET status = :stored_status WHERE status = :actual_status'),
                {"stored_status": "stored", "actual_status": "actual"}
            )
            updated_count = result.rowcount
            db.commit()  # Применяем обновление перед добавлением новых чанков
            if updated_count > 0:
                logger.info(f"Обновлено статусов существующих чанков на 'stored': {updated_count}")
            else:
                logger.debug("Нет чанков со статусом 'actual' для обновления")
            
            for chunk in chunks:
                try:
                    # Извлекаем метаданные
                    metadata = chunk.metadata or {}
                    chunk_metadata = {
                        k: v for k, v in metadata.items() 
                        if k not in ['chunk_id', 'total_chunks']
                    }
                    
                    db_chunk = Chunk(
                        content=chunk.page_content,
                        document_title=metadata.get("title"),
                        file_path=metadata.get("file_path"),
                        file_hash=metadata.get("file_hash"),
                        chunk_index=metadata.get("chunk_id"),
                        total_chunks=metadata.get("total_chunks"),
                        status="actual",  # Новые чанки получают статус "actual"
                        metadata_json=json.dumps(chunk_metadata, ensure_ascii=False) if chunk_metadata else None
                    )
                    db.add(db_chunk)
                    db.flush()  # Получаем ID без коммита
                    chunk_ids.append(db_chunk.id)
                except Exception as e:
                    logger.error(f"Ошибка при сохранении чанка в БД: {e}")
                    continue
            db.commit()
            logger.info(f"Сохранено {len(chunk_ids)} чанков в БД")
    except Exception as e:
        logger.error(f"Ошибка при сохранении чанков в БД: {e}")
    
    return chunk_ids


def split_documents(documents: List[Dict[str, Any]], 
                   chunk_size: int = None, 
                   chunk_overlap: int = None,
                   save_to_db: bool = True) -> List[LangChainDocument]:
    """
    Разбивает документы на чанки с помощью RecursiveCharacterTextSplitter
    и сохраняет их в базу данных
    
    Args:
        documents: Список документов для разбивки
        chunk_size: Размер чанка
        chunk_overlap: Перекрытие между чанками
        save_to_db: Сохранять ли чанки в БД (по умолчанию True)
    """
    if chunk_size is None:
        chunk_size = settings.chunk_size
    if chunk_overlap is None:
        chunk_overlap = settings.chunk_overlap
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""]
    )
    
    langchain_docs = []
    
    for i, doc in enumerate(documents, 1):
        try:
            langchain_doc = LangChainDocument(
                page_content=doc["content"],
                metadata={
                    "title": doc.get("title"),
                    "file_path": doc.get("file_path"),
                    "file_size": doc.get("file_size"),
                    "file_hash": doc.get("file_hash"),
                    "source": doc.get("title")
                }
            )

            chunks = text_splitter.split_documents([langchain_doc])
            
            langchain_docs.extend(chunks)
            logger.info(f"Документ '{doc['title']}' разбит на {len(chunks)} чанков")
            
        except Exception as e:
            logger.error(f"Ошибка при разбивке документа '{doc['title']}': {e}")
                
    # Добавляем глобальную нумерацию чанков (по всем документам)
    # chunk_index будет от 1 до общего количества всех чанков
    total_chunks = len(langchain_docs)
    for i, chunk in enumerate(langchain_docs, 1):
        chunk.metadata.update({
            "chunk_id": i,  # Глобальный индекс чанка (от 1 до total_chunks)
            "total_chunks": total_chunks  # Общее количество всех чанков
        })
    logger.info(f"Всего создано чанков: {total_chunks}, индексация: 1..{total_chunks}")
    
    # Сохраняем чанки в БД если нужно
    if save_to_db and langchain_docs:
        try:
            chunk_ids = save_chunks_to_db(langchain_docs)
            # Сохраняем ID в метаданные чанков для последующего использования
            for i, chunk in enumerate(langchain_docs):
                if i < len(chunk_ids):
                    chunk.metadata["db_id"] = chunk_ids[i]
        except Exception as e:
            logger.error(f"Ошибка при сохранении чанков в БД: {e}")
    
    return langchain_docs


def get_document_stats(documents: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Возвращает статистику по документам
    """
    if not documents:
        return {
            "total_documents": 0,
            "total_characters": 0,
            "total_size_bytes": 0,
            "average_document_size": 0
        }
    
    total_chars = sum(len(doc["content"]) for doc in documents)
    total_size = sum(doc["file_size"] for doc in documents)
    
    return {
        "total_documents": len(documents),
        "total_characters": total_chars,
        "total_size_bytes": total_size,
        "average_document_size": total_size / len(documents)
    } 